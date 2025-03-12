# 导入 python-docx 库中的 Document 类，用于处理 .docx 文件
from docx import Document
from tencent_client import batch_translate as tbt
from volcano_client import batch_translate as vbt
import re
from docx.shared import RGBColor
import json


def recognized_cn_en(input_text: str):
    # 定义正则表达式模式，用于匹配中文和英文部分
    # 匹配中文、中文标点 / 以及空格
    chinese_pattern = r'[\u4e00-\u9fa5()，。、；;：“”‘’（）—/ ]+'
    english_pattern = r'[a-zA-Z_\-, ;]+'
    # 使用正则表达式查找所有中文和英文部分
    chinese_list = re.findall(chinese_pattern, input_text)
    english_list = re.findall(english_pattern, input_text)
    # 去除 list 中的空串
    chinese_list = [chinese.replace(" ", "").strip() for chinese in chinese_list if chinese.strip()]
    english_list = [re.sub(r'\s+', ' ', eng).strip() for eng in english_list if eng.strip()]

    # 输出对齐的语料
    aligned_corpus = list(zip(chinese_list, english_list))
    warning_flag = False
    result = []
    warns = []
    for chinese, english in aligned_corpus:
        if len(chinese_list) != len(english_list):
            warning_flag = True
            # print(f'warning start ,raw text={input_text}')
        # print(f"{chinese.strip()}={english.strip()}")
        result.append(f'{chinese.strip()}={english.strip()}')
    if warning_flag:
        warns.append(f'{input_text}')
    return result, warns


# 定义函数 read_columns_from_docx，传入参数为 docx 文件路径
def extract_paragraphs_txt(docx_path):
    """
    从给定的docx文件中提取段落。

    参数:
    docx_path (str): docx文件的路径。

    返回:
    list: 包含文档中所有段落的列表，不包含空行。
    """
    # 加载文档
    doc = Document(docx_path)
    # 初始化段落列表
    paragraph_txt_list = []
    paragraph_list = []

    # 遍历文档中的每个段落
    for paragraph in doc.paragraphs:
        # 提取段落文本（自动过滤空行）
        text = paragraph.text.strip()
        # 仅保留非空内容
        if text:
            paragraph_txt_list.append(text)
            paragraph_list.append(paragraph)

    # 返回提取的段落列表
    return paragraph_txt_list


def contains_chinese(text):
    pattern = r'[\u4e00-\u9fa5()，。、；：“”‘’（）—/]+'
    return re.findall(pattern, text)


def read_vocabulary_from_docx(docx_path):
    # 打开并加载指定路径的 .docx 文件
    doc = Document(docx_path)
    result_list = []
    warn_list = []
    # 遍历文档中的所有段落
    for paragraph in doc.paragraphs:
        # 如果段落有非空文本，则将其添加到当前栏
        if not paragraph.text.strip():
            continue

        result, warns = recognized_cn_en(paragraph.text)
        result_list.extend(result)
        warn_list.extend(warns)
    with open('vocabulary.txt', 'w+', encoding='utf-8') as file:
        for line in result_list:
            file.write(line + '\n')

    with open('warns.txt', 'w+', encoding='utf-8') as file:
        for line in warn_list:
            file.write(line + '\n')


def extract_paragraphs(doc, en_map):
    # 初始化段落列表
    paragraph_list = []

    # 遍历文档中的每个段落
    for paragraph in doc.paragraphs:
        # 提取段落文本（自动过滤空行）
        text = paragraph.text.strip()
        # 仅保留非空内容
        if text:
            paragraph_list.append(paragraph)
    paragraph_list = paragraph_list[3:]
    idx = 0
    txt_set = set()
    while idx < len(paragraph_list):
        paragraph = paragraph_list[idx]
        text_strip = paragraph.text.strip()
        if contains_chinese(text_strip):
            txt = f'{text_strip}={paragraph_list[idx + 1].text.strip()}'
            txt_set.add(text_strip)
            if text_strip in en_map:
                run = paragraph_list[idx + 1].add_run(f'\n{en_map[text_strip]}')
                font = run.font
                font.color.rgb = RGBColor(255, 0, 0)
            idx += 2
        else:
            idx += 1
    # 返回提取的段落列表
    return txt_set


def extract_all_tables(doc, en_map):
    tables = doc.tables
    txt_set = set()
    for idx, table in enumerate(tables):
        cn = []
        en = []
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                # 检查单元格是否是合并单元格的一部分
                paragraphs = [para.text.strip().replace('..', '').replace('\n', '') for para in cell.paragraphs if
                              para.text.strip()]
                if len(paragraphs) == 1:
                    continue
                elif len(paragraphs) > 2 and len(paragraphs) & 1 == 0:
                    idx = 0
                    while idx < len(paragraphs):
                        paragraph = paragraphs[idx]
                        if contains_chinese(paragraph):
                            cn.append(paragraph)
                            en.append(paragraphs[idx + 1])
                            if paragraph in en_map:
                                run = cell.paragraphs[idx + 1].add_run(f'\n{en_map[paragraph]}')
                                font = run.font
                                font.color.rgb = RGBColor(255, 0, 0)
                            idx += 2
                        else:
                            idx += 1
                elif len(paragraphs) == 2:
                    cn.append(paragraphs[0])
                    en.append(paragraphs[1])
                    if paragraphs[0] in en_map:
                        run = cell.paragraphs[1].add_run(f'\n{en_map[paragraphs[0]]}')
                        font = run.font
                        font.color.rgb = RGBColor(255, 0, 0)
            result = list(zip(cn, en))
            for c, e in result:
                txt = f'{c}={e}'
                txt_set.add(c)
    return txt_set


def cn_to_en(cn_text):
    pass


def translate_in_batches(union_set, output_file):
    batches = [list(union_set)[i:i + 10] for i in range(0, len(union_set), 10)]
    translation_map = {}
    for batch in batches:
        translated_texts = tbt(batch)
        for original, translated in zip(batch, translated_texts):
            translation_map[original] = translated
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(translation_map, f, ensure_ascii=False, indent=4)


if __name__ == '__main__':
    with open('test1.json', 'r', encoding='utf-8') as file:
        en_map = json.load(file)
        doc = Document('/Users/deipss/Desktop/LZTD/a供电专业课程国际化翻译/test1.docx')
        txt_set = extract_paragraphs(doc,en_map)
        txt_set_table = extract_all_tables(doc,en_map)
        doc.save('z.docx')